import sys
sys.path.insert(0, '/opt/homebrew/lib/python3.9/site-packages')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime
import json, subprocess, urllib.request

def gerar_proposta(cliente, sindico, unidades, servicos, valor_mensal, implantacao):
    doc = Document()

    titulo = doc.add_heading('', 0)
    run = titulo.add_run('WPS DIGITAL')
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x9b)
    run.font.size = Pt(28)

    sub = doc.add_paragraph()
    run2 = sub.add_run('Proposta Comercial de Segurança e Tecnologia')
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    info = doc.add_table(rows=4, cols=2)
    info.style = 'Table Grid'
    dados = [
        ('Cliente', cliente),
        ('Síndico', sindico),
        ('Unidades', str(unidades)),
        ('Data', datetime.now().strftime('%d/%m/%Y')),
    ]
    for i, (k, v) in enumerate(dados):
        info.rows[i].cells[0].text = k
        info.rows[i].cells[1].text = v

    doc.add_paragraph()
    doc.add_heading('Diagnóstico', level=1)
    doc.add_paragraph(f'Condomínio com {unidades} unidades necessita de solução integrada de segurança, '
        f'controle de acesso e conectividade. A WPS Digital, com 25 anos de experiência e '
        f'500+ clientes, oferece a solução ideal para este perfil.')

    doc.add_heading('Solução Proposta', level=1)
    for s in servicos:
        doc.add_paragraph(f'• {s}', style='Normal')

    doc.add_heading('Investimento', level=1)
    inv = doc.add_table(rows=3, cols=2)
    inv.style = 'Table Grid'
    economia = int(valor_mensal * 1.7)
    roi = round(implantacao / max(economia - valor_mensal, 1), 1)
    investimentos = [
        ('Implantação', f'R$ {implantacao:,.0f}'),
        ('Mensalidade', f'R$ {valor_mensal:,.0f}/mês'),
        ('ROI estimado', f'{roi} meses'),
    ]
    for i, (k, v) in enumerate(investimentos):
        inv.rows[i].cells[0].text = k
        inv.rows[i].cells[1].text = v

    doc.add_paragraph()
    rodape = doc.add_paragraph()
    run3 = rodape.add_run('WPS Digital — 25 anos | 500+ condomínios | suporte@wpsdigital.com.br')
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    filename = f'/tmp/proposta_wps_{cliente.replace(" ","_")}.docx'
    doc.save(filename)
    return filename

if __name__ == '__main__':
    f = gerar_proposta(
        cliente='Condomínio São Francisco',
        sindico='Sr. João Silva',
        unidades=280,
        servicos=['CFTV Hikvision 32 câmeras HD', 'Portaria Virtual 24h', 'Rede WiFi áreas comuns', 'Controle de Acesso Biométrico'],
        valor_mensal=5200,
        implantacao=42000
    )
    print(f'Proposta gerada: {f}')
